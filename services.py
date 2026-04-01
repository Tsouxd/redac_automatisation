import os
import google.generativeai as genai
from docx import Document
import shutil

# --- CONFIGURATION GEMINI ---
# On récupère la clé depuis les variables d'environnement du serveur
API_KEY = os.environ.get("GEMINI_API_KEY")

if not API_KEY:
    print("ERREUR : La variable d'environnement GEMINI_API_KEY n'est pas configurée.")
else:
    genai.configure(api_key=API_KEY)
    
# Configuration de l'API
genai.configure(api_key=API_KEY)

# Votre liste de priorité
GEMINI_MODEL_FALLBACK_LIST = [
    "gemini-3-flash-preview",
    "gemini-2.5-flash-lite",
    "gemini-2.5-flash"
]

def extract_text_from_file(filepath):
    """Lit le fichier (.txt ou .docx) et retourne le texte."""
    ext = filepath.split('.')[-1].lower()
    text = ""
    try:
        if ext == 'txt':
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        elif ext == 'docx':
            doc = Document(filepath)
            text = '\n'.join([para.text for para in doc.paragraphs])
        else:
            return None
    except Exception:
        return None
    return text

def process_with_gemini(source_text, example_text):
    """
    Appelle Gemini avec une logique de fallback sur les modèles.
    """
    
    # Construction du prompt
    prompt = f"""
    Tu es un assistant rédactionnel expert.
    
    TÂCHE : Réécris le 'CONTENU À TRAITER' en imitant parfaitement le style, le ton et la structure du 'MODÈLE DE STYLE'.
    
    --- DÉBUT MODÈLE DE STYLE ---
    {example_text[:5000]} 
    --- FIN MODÈLE DE STYLE ---
    
    --- DÉBUT CONTENU À TRAITER ---
    {source_text}
    --- FIN CONTENU À TRAITER ---
    
    Génère uniquement le document réécrit final, sans phrase d'introduction ni de conclusion hors contexte.
    """

    last_error = None

    # Boucle sur la liste des modèles (Fallback)
    for model_name in GEMINI_MODEL_FALLBACK_LIST:
        try:
            print(f"Tentative avec le modèle : {model_name}...")
            model = genai.GenerativeModel(model_name)
            
            # Appel à l'API
            response = model.generate_content(prompt)
            
            # Si succès, on retourne le texte et on sort de la boucle
            print(f"Succès avec {model_name} !")
            return response.text
            
        except Exception as e:
            print(f"Échec avec {model_name}. Erreur : {e}")
            last_error = e
            continue # On passe au modèle suivant dans la liste

    # Si tous les modèles échouent
    return f"ERREUR : Impossible de générer le texte avec les modèles fournis. Détails : {last_error}"

def generate_word_doc_from_template(ai_content, folder_resultats, path_modele_original):
    """
    1. Copie le modèle original.
    2. Vide le texte de la copie.
    3. Remplit avec le texte de l'IA.
    """
    
    # Nom du futur fichier final
    filename_final = "Resultat_Final.docx"
    path_final = os.path.join(folder_resultats, filename_final)
    
    # ÉTAPE 1 : VÉRIFICATION ET COPIE SÉCURISÉE
    # Si le modèle n'est pas un .docx, on ne peut pas cloner le style -> on crée un nouveau doc
    if not path_modele_original.endswith('.docx'):
        doc = Document()
        doc.add_paragraph(ai_content)
        doc.save(path_final)
        return filename_final

    # On DUPLIQUE le fichier modèle vers le dossier de résultats
    shutil.copy(path_modele_original, path_final)
    
    try:
        # ÉTAPE 2 : OUVERTURE DE LA COPIE
        doc = Document(path_final)
        
        # ÉTAPE 3 : VIDER LE CONTENU (CORPS DU TEXTE)
        # On supprime tous les paragraphes existants pour faire place nette
        # (Les en-têtes, pieds de page et marges restent intacts !)
        for element in doc.element.body:
            doc.element.body.remove(element)
            
        # ÉTAPE 4 : INJECTION DU TEXTE IA
        # On ajoute le texte. Il prendra automatiquement le style "Normal" du document modèle.
        
        # Titre (Optionnel : on ajoute un titre pour faire propre)
        doc.add_heading('Document Généré', 0)

        for line in ai_content.split('\n'):
            line = line.strip()
            if line:
                # Détection simple de titre (Ligne courte et Majuscules)
                if len(line) < 60 and line.isupper():
                    # On tente d'utiliser le style "Heading 1" du modèle s'il existe
                    try:
                        doc.add_heading(line, level=1)
                    except:
                        doc.add_paragraph(line) # Repli si le style n'existe pas
                else:
                    doc.add_paragraph(line)

        # ÉTAPE 5 : SAUVEGARDE SUR LA COPIE
        doc.save(path_final)
        return filename_final

    except Exception as e:
        print(f"Erreur lors de la manipulation du template : {e}")
        # En cas de crash, on rend un fichier simple sans style
        doc = Document()
        doc.add_paragraph(ai_content)
        doc.save(path_final)
        return filename_final